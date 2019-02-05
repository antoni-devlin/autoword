from docx import Document
from docx.shared import Inches
import csv

document = Document()

csv = input("Enter name of CSV file (music.csv): ")

with open(csv, newline='') as f:

	csv_reader = csv.reader(f)

	csv_headers = next(csv_reader)
	csv_cols = len(csv_headers)

	for row in csv_reader:
		document.add_paragraph(row[0])
		document.add_paragraph(row[1])
		document.add_paragraph(row[2])

		document.save(row[1]+'.docx')

